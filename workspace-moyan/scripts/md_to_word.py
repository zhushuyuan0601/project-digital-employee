#!/usr/bin/env python3
"""周报 Markdown 转 Word 脚本"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys

def parse_markdown_to_word(md_file, docx_file):
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)  # 五号字
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 标题
        if line.startswith('# 周报'):
            p = doc.add_heading('周报', 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # 时间范围和汇报人
        if line.startswith('**时间范围**'):
            # 提取时间范围
            match = re.search(r'\*时间范围\*\*: (.+?)\n\*\*汇报人\*\*: (.+?)\n', content)
            if match:
                time_range = match.group(1)
                reporter = match.group(2)
                p = doc.add_paragraph()
                p.add_run(f'时间范围：{time_range}\n')
                p.add_run(f'汇报人：{reporter}')
                p.paragraph_format.space_after = Pt(12)
            i += 2
            continue
        
        # 一级标题（## 一、xxx）
        if line.startswith('## 一、') or line.startswith('## 二、') or line.startswith('## 三、') or \
           line.startswith('## 四、') or line.startswith('## 五、') or line.startswith('## 六、'):
            title = line.replace('## ', '')
            # 如果包含"本周无相关工作"，用较小的字体
            if i + 1 < len(lines) and '本周无相关工作' in lines[i + 1]:
                p = doc.add_heading(title, level=1)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                i += 1
                if i < len(lines) and '本周无相关工作' in lines[i]:
                    doc.add_paragraph('本周无相关工作。')
                    i += 1
                continue
            else:
                p = doc.add_heading(title, level=1)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                i += 1
                continue
        
        # 二级标题（### 1. xxx）
        if line.startswith('### '):
            title = line.replace('### ', '')
            p = doc.add_heading(title, level=2)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue
        
        # 粗体标签
        if line.startswith('**目标**'):
            match = re.match(r'\*\*(.+?)\*\*: (.+)', line)
            if match:
                label = match.group(1).replace('*', '')
                value = match.group(2)
                p = doc.add_paragraph()
                p.add_run(f'{label}：').bold = True
                p.add_run(value)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(0)
                i += 1
                continue
        
        if line.startswith('**本周进展**'):
            match = re.match(r'\*\*(.+?)\*\*: (.+)', line)
            if match:
                label = match.group(1).replace('*', '')
                value = match.group(2)
                p = doc.add_paragraph()
                p.add_run(f'{label}：').bold = True
                p.add_run(value)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(0)
                i += 1
                continue
        
        if line.startswith('**完成度**'):
            match = re.match(r'\*\*(.+?)\*\*: (.+)', line)
            if match:
                label = match.group(1).replace('*', '')
                value = match.group(2)
                p = doc.add_paragraph()
                p.add_run(f'{label}：').bold = True
                p.add_run(value)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(0)
                i += 1
                continue
        
        if line.startswith('**协同人员**'):
            match = re.match(r'\*\*(.+?)\*\*: (.+)', line)
            if match:
                label = match.group(1).replace('*', '')
                value = match.group(2)
                p = doc.add_paragraph()
                p.add_run(f'{label}：').bold = True
                p.add_run(value)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(0)
                i += 1
                continue
        
        if line.startswith('**预计完成时间**'):
            match = re.match(r'\*\*(.+?)\*\*: (.+)', line)
            if match:
                label = match.group(1).replace('*', '')
                value = match.group(2)
                p = doc.add_paragraph()
                p.add_run(f'{label}：').bold = True
                p.add_run(value)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(0)
                i += 1
                continue
        
        if line.startswith('**下一步**'):
            match = re.match(r'\*\*(.+?)\*\*: (.+)', line)
            if match:
                label = match.group(1).replace('*', '')
                value = match.group(2)
                p = doc.add_paragraph()
                p.add_run(f'{label}：').bold = True
                p.add_run(value)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(0)
                i += 1
                continue
        
        # 列表项
        if line.startswith('- '):
            content = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(content)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue
        
        # 分割线
        if line.startswith('---'):
            i += 1
            continue
        
        # 备注部分
        if line.startswith('## 备注'):
            p = doc.add_heading('备注', level=1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        if line.startswith('## 下周计划'):
            p = doc.add_heading('下周计划', level=1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # 普通段落
        if line.strip() and not line.startswith('```'):
            p = doc.add_paragraph(line)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue
        
        i += 1
    
    doc.save(docx_file)
    print(f'✅ Word 文档已生成：{docx_file}')

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('用法：python3 md_to_word.py <markdown 文件> <word 文件>')
        sys.exit(1)
    
    parse_markdown_to_word(sys.argv[1], sys.argv[2])
